import random
import math
from docx import Document
from docx.shared import Pt
import os
import sys

quotes = []
# font_list = ['Arial', 'Times New Roman', 'Helvetica', 'Calibri', 'Verdana', 'Georgia', 'Courier New', 'Tahoma']
font_list = [
    'Arial', 'Times New Roman', 'Helvetica', 'Calibri', 'Verdana', 'Georgia',
    'Courier New', 'Tahoma', 'Trebuchet MS', 'Lucida Sans', 'Garamond', 'Palatino Linotype',
    'Consolas', 'Franklin Gothic Medium', 'Gill Sans', 'Cambria', 'Segoe UI',
    'Candara', 'Perpetua', 'Book Antiqua', 'Corbel', 'Century Gothic', 'Rockwell',
    'Didot', 'Baskerville', 'Futura', 'Optima', 'American Typewriter', 'Charter',
    'Lucida Bright', 'Monaco', 'Courier', 'Lucida Console', 'Microsoft Sans Serif',
    'Berlin Sans FB', 'Century Schoolbook', 'Bookman Old Style', 'Tw Cen MT',
    'Tempus Sans ITC', 'Ebrima', 'Constantia', 'Bahnschrift', 'Leelawadee UI',
    'Sitka Text', 'Utopia', 'Mongolian Baiti', 'Noto Sans', 'PT Sans', 'Quicksand', 
    'Roboto', 'Open Sans', 'Lato'
]

doc = Document()

def add_letter_with_random_font(paragraph, letter):
    run = paragraph.add_run(letter)
    font_name = random.choice(font_list)
    run.font.name = font_name
    size = math.ceil(random.randint(1,3)) + 7
    run.font.size = Pt(size)  # Adjust the font size if needed

with open('quotes.txt', 'r') as file:
    lines = file.readlines()
for line in lines:
    line = line.strip()
    quotes.append(line)

new_combo = []

sorted_quotes = sorted(quotes, key=len)

i = 0
while i < len(sorted_quotes) / 2:
    new_combo.append([sorted_quotes[i], sorted_quotes[len(sorted_quotes) - i - 1]])
    i+=1

final_list = []

i = 0
while i < len(new_combo):
    if i < len(new_combo) / 2 * 999999:
        if i % 2 == 0:
            final_list.append(new_combo[i][0] + " // " + new_combo[i][1])
        else:
            final_list.append(new_combo[i][0] + " \\\ " + new_combo[i][1])
    else:
        if i % 2 == 0:
            final_list.append(new_combo[i][1] + " // " + new_combo[i][0])
        else:
            final_list.append(new_combo[i][1] + " \\\ " + new_combo[i][0])
    i+=1

for string in final_list:
    paragraph = doc.add_paragraph()
    for letter in string:
        add_letter_with_random_font(paragraph, letter)

doc.save('randomized_fonts.docx')